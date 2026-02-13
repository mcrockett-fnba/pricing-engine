#!/usr/bin/env python3
"""Generate a management-facing Word document summarizing the APEX2 review."""
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT_DIR = Path(__file__).resolve().parent.parent.parent / "reports"
OUT_DIR.mkdir(exist_ok=True)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    return h


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(f"  {text}")
    else:
        p.add_run(text)
    return p


def main():
    doc = Document()

    # -- Title --
    title = doc.add_heading("APEX2 Prepayment Model Review", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    subtitle = doc.add_paragraph("Executive Summary for Leadership")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    doc.add_paragraph("")  # spacer

    # -- Executive Summary --
    add_heading(doc, "Executive Summary")
    doc.add_paragraph(
        "APEX2 is our production prepayment model used to estimate how quickly "
        "borrowers pay ahead of schedule, which directly drives loan pricing and "
        "return projections. This review assessed the model's strengths and "
        "identified targeted enhancements for the next-generation pricing engine."
    )
    doc.add_paragraph(
        "The core finding is positive: APEX2 is built on a solid, data-driven "
        "foundation that correctly identifies the key factors influencing "
        "prepayment behavior. The recommended improvements focus on refining "
        "how those factors are combined, not replacing them."
    )

    # -- What APEX2 Gets Right --
    add_heading(doc, "Current Strengths")
    doc.add_paragraph(
        "APEX2 represents a significant advancement over the prior system (PPD_OLD), "
        "which had no analytical pricing layer. APEX2 introduced data-driven "
        "prepayment estimation, ROE-based pricing, and multi-dimensional analysis "
        "across key loan characteristics."
    )

    add_bullet(doc,
        "Multipliers are grounded in actual loan performance data from our "
        "portfolio, not industry assumptions or generic benchmarks.",
        "Data-Driven Foundation.")

    add_bullet(doc,
        "The model evaluates loans across multiple characteristics simultaneously "
        "\u2014 credit score, rate environment, LTV, loan size, and more \u2014 giving "
        "a nuanced view of each loan\u2019s prepayment profile.",
        "Multi-Dimensional Analysis.")

    add_bullet(doc,
        "APEX2 recognizes that ITIN and non-ITIN borrowers exhibit fundamentally "
        "different prepayment patterns and maintains separate calibrations for each. "
        "Our ITIN dataset is likely one of the best in the industry for this segment.",
        "ITIN Stratification.")

    add_bullet(doc,
        "Prepayment estimates flow directly into ROE-based bid pricing, creating "
        "a tight link between model output and business decisions.",
        "Integrated Pricing Pipeline.")

    add_bullet(doc,
        "The dimensional lookup approach is straightforward enough that analysts "
        "can understand, explain, and validate results \u2014 an important attribute "
        "for a model that drives pricing decisions.",
        "Transparency and Auditability.")

    # -- Key Dimensions --
    add_heading(doc, "How APEX2 Works")
    doc.add_paragraph(
        "APEX2 assigns each loan a prepayment speed multiplier by looking up "
        "the loan\u2019s characteristics across dimensional tables. Key dimensions include:"
    )

    add_bullet(doc,
        "Higher credit scores correlate with faster prepayment (2.7x for 750+ vs. "
        "1.4x for sub-576). This is the strongest single predictor.",
        "Credit Score.")

    add_bullet(doc,
        "Loans with rates well above the current market prepay faster (borrowers "
        "have refinance incentive). Loans below market prepay slowly.",
        "Rate vs. Market.")

    add_bullet(doc,
        "Lower LTV borrowers (more equity) prepay faster \u2014 they have more "
        "refinance options. High-LTV borrowers are more constrained.",
        "Loan-to-Value.")

    add_bullet(doc,
        "Larger loans prepay materially faster than smaller loans (3.3x for $1M+ "
        "vs. 1.3x for under $50K), reflecting stronger financial sophistication "
        "and refinance economics.",
        "Loan Size.")

    doc.add_paragraph(
        "These multipliers determine how quickly each loan is expected to pay off "
        "relative to its scheduled amortization, which directly impacts projected "
        "returns and bid pricing."
    )

    # -- What's New in the Engine --
    add_heading(doc, "Enhancements Already Implemented")
    doc.add_paragraph(
        "The new pricing engine has already incorporated several improvements "
        "to the APEX2 framework:"
    )

    add_bullet(doc,
        "The engine now applies a 30-month seasoning ramp to new loans, "
        "reflecting the industry-standard observation that borrowers rarely "
        "refinance in the first 1\u20132 years. This improves accuracy for "
        "packages containing a mix of new and seasoned loans.",
        "Seasoning Adjustment.")

    add_bullet(doc,
        "Analysts can now compare results across six scenarios (two multiplier "
        "sources \u00d7 three projection methods), providing better insight into "
        "how sensitive a package\u2019s effective life is to modeling assumptions.",
        "Multiple Scenario Analysis.")

    add_bullet(doc,
        "The APEX2 lookup tables are now registered as a versioned model "
        "in the engine\u2019s model registry, enabling version tracking, "
        "auditability, and the ability to swap in updated tables without "
        "code changes.",
        "Model Registry Integration.")

    add_bullet(doc,
        "Analysts can upload loan tapes directly from Excel, replacing "
        "the previous manual data entry process. Column matching is flexible "
        "and unit conversions (rates, LTV) are handled automatically.",
        "Excel Tape Upload.")

    # -- Roadmap --
    add_heading(doc, "Recommended Next Steps")
    doc.add_paragraph(
        "The review identified several areas where targeted improvements could "
        "enhance model accuracy. These build on the existing APEX2 foundation "
        "rather than replacing it."
    )

    add_bullet(doc,
        "Use regression or machine learning to determine how much weight each "
        "dimension should carry, rather than weighting all equally. Credit score "
        "and rate environment are likely more predictive than some other factors.",
        "Optimized Dimension Weighting.")

    add_bullet(doc,
        "Track how actual loan runoff compares to APEX2 predictions, "
        "particularly for non-ITIN packages where our internal data is smaller "
        "relative to the broader market.",
        "Post-Acquisition Performance Tracking.")

    add_bullet(doc,
        "More recent loan performance data should carry greater weight than "
        "observations from 2007\u20132015, as prepayment behavior varies "
        "significantly across different rate environments.",
        "Recency Weighting.")

    add_bullet(doc,
        "Given the meaningful differences in data quality and borrower behavior "
        "between ITIN and non-ITIN segments, maintaining distinct calibration "
        "paths will improve accuracy for both.",
        "Separate ITIN / Non-ITIN Calibration.")

    # -- Closing --
    add_heading(doc, "Conclusion")
    doc.add_paragraph(
        "APEX2 provides a solid, proven foundation for prepayment estimation. "
        "It correctly identifies the key loan characteristics that drive prepayment "
        "behavior and is grounded in real portfolio data. The new pricing engine "
        "preserves these strengths while adding seasoning adjustments, scenario "
        "analysis, model versioning, and streamlined data ingestion. The recommended "
        "enhancements \u2014 optimized weighting, performance tracking, and recency "
        "adjustments \u2014 represent incremental improvements that can be phased in "
        "over time without disrupting current operations."
    )

    # -- Save --
    out_path = OUT_DIR / "APEX2_Review_Summary.docx"
    doc.save(str(out_path))
    print(f"Saved to {out_path}")


if __name__ == "__main__":
    main()
